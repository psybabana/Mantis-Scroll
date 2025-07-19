import tkinter as tk
from tkinter import filedialog, messagebox, font
from tkinter.scrolledtext import ScrolledText
from docx import Document
from PIL import Image, ImageTk, ImageEnhance
import time
import threading

class WordProcessorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Mantis Scroll")
        self.root.attributes("-fullscreen", True)
        self.filename = None

        self.bg_image_path = None
        self.bg_opacity = 0.3  # default opacity (30%)

        # Timer state
        self.timer_running = False
        self.start_time = None
        self.elapsed_time = 0
        self.toolbar = tk.Frame(self.root)
        self.toolbar.pack(side=tk.TOP, fill=tk.X)

        self.create_menu()
        self.create_timer_toolbar()
        self.create_formatting_toolbar()
        self.create_text_area()

        self.root.bind("<Escape>", self.exit_fullscreen)

    def create_menu(self):
        menubar = tk.Menu(self.root)

        #File Menu
        file_menu = tk.Menu(menubar, tearoff=0)
        file_menu.add_command(label="Open", command=self.open_file)
        file_menu.add_command(label="Save", command=self.save_file)
        file_menu.add_command(label="Save As", command=self.save_as)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self.root.quit)
        menubar.add_cascade(label="File", menu=file_menu)

        #Theme Menu
        theme_menu = tk.Menu(menubar, tearoff=0)
        theme_menu.add_command(label="Set Inspiration Image", command=self.select_background_image)
        menubar.add_cascade(label="Theme", menu=theme_menu)

        #Window Menu
        window_menu = tk.Menu(menubar, tearoff=0)
        window_menu.add_command(label="Toggle Fullscreen", command=self.toggle_fullscreen)
        menubar.add_cascade(label="Window", menu=window_menu)

        self.root.config(menu=menubar)

    def create_timer_toolbar(self):
        timer_frame = tk.Frame(self.root, relief=tk.RAISED, bd=1)
        timer_frame.pack(side=tk.TOP, fill=tk.X)
        
        self.timer_text_label = tk.Label(timer_frame, text="Timer:", font=("Arial", 12, "bold"))
        self.timer_text_label.pack(side=tk.LEFT, padx=(10, 2))

        self.timer_label = tk.Label(timer_frame, text="00:00:00.000", font=("Courier", 12, "bold"))
        self.timer_label.pack(side=tk.LEFT, padx=10)

        tk.Button(timer_frame, text="Start", command=self.start_timer).pack(side=tk.LEFT)
        tk.Button(timer_frame, text="Pause", command=self.pause_timer).pack(side=tk.LEFT)
        tk.Button(timer_frame, text="Reset", command=self.reset_timer).pack(side=tk.LEFT)

        self.update_timer_loop()

    def start_timer(self):
        if not self.timer_running:
            self.start_time = time.time() - self.elapsed_time
            self.timer_running = True

    def pause_timer(self):
        if self.timer_running:
            self.elapsed_time = time.time() - self.start_time
            self.timer_running = False

    def reset_timer(self):
        self.timer_running = False
        self.elapsed_time = 0
        self.timer_label.config(text="00:00:00.000")

    def update_timer_loop(self):
        if self.timer_running:
            now = time.time()
            elapsed = now - self.start_time
            self.elapsed_time = elapsed

            ms = int((elapsed - int(elapsed)) * 1000)
            s = int(elapsed) % 60
            m = (int(elapsed) // 60) % 60
            h = int(elapsed) // 3600

            self.timer_label.config(text=f"{h:02d}:{m:02d}:{s:02d}.{ms:03d}")
        self.root.after(50, self.update_timer_loop)

    def toggle_fullscreen(self):
        is_fullscreen = self.root.attributes("-fullscreen")
        self.root.attributes("-fullscreen", not is_fullscreen)

    def create_formatting_toolbar(self):
        toolbar = tk.Frame(self.root, relief=tk.RAISED, bd=1)
        toolbar.pack(side=tk.TOP, fill=tk.X)

        self.bold_btn = tk.Button(toolbar, text="B", width=2, font=("Arial", 10, "bold"), command=self.toggle_bold)
        self.bold_btn.pack(side=tk.LEFT, padx=2, pady=2)

        self.italic_btn = tk.Button(toolbar, text="I", width=2, font=("Arial", 10, "italic"), command=self.toggle_italic)
        self.italic_btn.pack(side=tk.LEFT, padx=2, pady=2)

        self.underline_btn = tk.Button(toolbar, text="U", width=2, font=("Arial", 10, "underline"), command=self.toggle_underline)
        self.underline_btn.pack(side=tk.LEFT, padx=2, pady=2)

        self.subscript_btn = tk.Button(toolbar, text="x₂", width=3, command=self.toggle_subscript)
        self.subscript_btn.pack(side=tk.LEFT, padx=2, pady=2)

        self.superscript_btn = tk.Button(toolbar, text="x²", width=3, command=self.toggle_superscript)
        self.superscript_btn.pack(side=tk.LEFT, padx=2, pady=2)

        self.font_family_var = tk.StringVar(value="Arial")
        font_families = sorted(list(font.families()))
        self.font_family_menu = tk.OptionMenu(toolbar, self.font_family_var, *font_families, command=self.change_font)
        self.font_family_menu.config(width=15)
        self.font_family_menu.pack(side=tk.LEFT, padx=5, pady=2)

        self.font_size_var = tk.IntVar(value=12)
        self.font_size_spinbox = tk.Spinbox(toolbar, from_=8, to=72, textvariable=self.font_size_var,
                                            width=5, command=self.change_font)
        self.font_size_spinbox.pack(side=tk.LEFT, padx=5, pady=2)

        tk.Label(toolbar, text="Inspiration Image Opacity").pack(side=tk.LEFT, padx=5)
        self.opacity_var = tk.DoubleVar(value=self.bg_opacity * 100)
        self.opacity_slider = tk.Scale(toolbar, from_=0, to=100, orient=tk.HORIZONTAL, length=100,
                                       variable=self.opacity_var, command=self.on_opacity_change)
        self.opacity_slider.pack(side=tk.LEFT, padx=5, pady=2)

        #Fullscreen+
        self.fullscreen_btn = tk.Button(self.toolbar, text="Fullscreen", command=self.toggle_fullscreen)
        self.fullscreen_btn.pack(side=tk.RIGHT, padx=2)


    def create_text_area(self):
        self.container = tk.Frame(self.root)
        self.container.pack(fill=tk.BOTH, expand=True)

         # Fixed-size frame for the background image
        self.bg_frame = tk.Frame(self.container, width=700)
        self.bg_frame.pack(side=tk.LEFT, fill=tk.Y)
        self.bg_frame.pack_propagate(False)

        self.bg_canvas = tk.Canvas(self.bg_frame, highlightthickness=10)
        self.bg_canvas.pack(fill=tk.BOTH, expand=True)

         # Expandable frame for text area
        self.text_area = ScrolledText(self.container, wrap=tk.WORD, undo=True,
                                 font=(self.font_family_var.get(), self.font_size_var.get()),
                                 bg="#ffffff", relief=tk.FLAT)
        self.text_area.pack(fill=tk.BOTH, expand=True, side=tk.LEFT)

        self.text_area.vbar.config(command=self.on_textscroll)
        self.bg_canvas.bind("<Configure>", self.resize_bg)
        self.bg_img_original = None
        self.bg_img = None

    def select_background_image(self):
        path = filedialog.askopenfilename(title="Select Inspiration Image",
                                          filetypes=[("Image files", "*.png *.jpg *.jpeg *.bmp *.gif")])
        if not path:
            return
        self.bg_image_path = path
        self.load_and_set_background(self.bg_image_path, self.bg_opacity)

    def load_and_set_background(self, path, opacity):
        try:
            img = Image.open(path).convert("RGBA")
            alpha = img.split()[3]
            alpha = ImageEnhance.Brightness(alpha).enhance(opacity)
            img.putalpha(alpha)
            self.bg_img_original = img
            self.update_bg_image()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load image:\n{e}")
            self.bg_img_original = None
            self.bg_img = None
            self.bg_canvas.delete("all")

    def update_bg_image(self):
        if not self.bg_img_original:
            self.bg_canvas.delete("all")
            return

        w = self.bg_canvas.winfo_width()
        h = self.bg_canvas.winfo_height()
        if w < 10 or h < 10:
            return

        resized = self.bg_img_original.resize((w, h), Image.LANCZOS)
        self.bg_img = ImageTk.PhotoImage(resized)
        self.bg_canvas.delete("all")
        self.bg_canvas.create_image(0, 0, anchor=tk.NW, image=self.bg_img)

    def resize_bg(self, event=None):
        self.update_bg_image()

    def on_opacity_change(self, value):
        opacity = float(value) / 100.0
        self.bg_opacity = opacity
        if self.bg_image_path:
            self.load_and_set_background(self.bg_image_path, opacity)

    def on_textscroll(self, *args):
        self.text_area.yview(*args)
        self.bg_canvas.yview_moveto(args[1])

    def exit_fullscreen(self, event=None):
        self.root.attributes("-fullscreen", False)

    def toggle_tag(self, tag_name, font_mod):
        try:
            start = self.text_area.index("sel.first")
            end = self.text_area.index("sel.last")
        except tk.TclError:
            return

        if tag_name in self.text_area.tag_names("sel.first"):
            self.text_area.tag_remove(tag_name, start, end)
        else:
            self.text_area.tag_add(tag_name, start, end)
            self.text_area.tag_configure(tag_name, font=font_mod)

    def toggle_bold(self):
        current_font = font.Font(self.text_area, self.text_area.cget("font"))
        bold_font = font.Font(family=current_font.actual("family"), size=current_font.actual("size"), weight="bold")
        self.toggle_tag("bold", bold_font)

    def toggle_italic(self):
        current_font = font.Font(self.text_area, self.text_area.cget("font"))
        italic_font = font.Font(family=current_font.actual("family"), size=current_font.actual("size"), slant="italic")
        self.toggle_tag("italic", italic_font)

    def toggle_underline(self):
        current_font = font.Font(self.text_area, self.text_area.cget("font"))
        underline_font = font.Font(family=current_font.actual("family"), size=current_font.actual("size"), underline=1)
        self.toggle_tag("underline", underline_font)

    def toggle_subscript(self):
        self.toggle_tag("subscript", font.Font(self.text_area, self.text_area.cget("font")))
        self.text_area.tag_configure("subscript", offset=-3)

    def toggle_superscript(self):
        self.toggle_tag("superscript", font.Font(self.text_area, self.text_area.cget("font")))
        self.text_area.tag_configure("superscript", offset=3)

    def change_font(self, *args):
        new_font = (self.font_family_var.get(), self.font_size_var.get())
        self.text_area.configure(font=new_font)

    def save_file(self):
        if self.filename:
            self._save(self.filename)
        else:
            self.save_as()

    def save_as(self):
        filetypes = [("Text files", ".txt"), ("Word documents", ".docx")]
        path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=filetypes)
        if not path:
            return
        self.filename = path
        if path.endswith(".txt"):
            messagebox.showwarning("Formatting lost", "Saving as .txt will not preserve formatting.")
        self._save(path)

    def _save(self, path):
        content = self.text_area.get("1.0", "end-1c")
        if path.endswith(".txt"):
            with open(path, "w", encoding="utf-8") as f:
                f.write(content)
        elif path.endswith(".docx"):
            doc = Document()
            for line in content.splitlines():
                doc.add_paragraph(line)
            doc.save(path)

    def open_file(self):
        filetypes = [("Text files", ".txt"), ("Word documents", ".docx")]
        path = filedialog.askopenfilename(filetypes=filetypes)
        if not path:
            return
        self.filename = path
        self.text_area.delete("1.0", "end")
        if path.endswith(".txt"):
            with open(path, "r", encoding="utf-8") as f:
                self.text_area.insert("1.0", f.read())
        elif path.endswith(".docx"):
            doc = Document(path)
            text = "\n".join(para.text for para in doc.paragraphs)
            self.text_area.insert("1.0", text)


if __name__ == "__main__":
    root = tk.Tk()
    app = WordProcessorApp(root)
    root.mainloop()
